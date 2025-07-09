
import streamlit as st
from docx import Document
from docx.shared import Inches
import io
from datetime import datetime

# Helper function
def romawi(bulan):
    mapping = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X", "XI", "XII"]
    return mapping[bulan - 1]

def generate_word(data, items):
    doc = Document("template.docx")

    # Ganti teks di dokumen berdasarkan data
    for p in doc.paragraphs:
        for key, value in data.items():
            if key in p.text:
                p.text = p.text.replace(key, value)

    # Tambah tabel item
    for table in doc.tables:
        if "{{TABLE}}" in table.cell(0,0).text:
            table.cell(0,0).text = "Qty"
            table.cell(0,1).text = "Part Number"
            table.cell(0,2).text = "Description"
            table.cell(0,3).text = "Price"
            for item in items:
                row = table.add_row().cells
                row[0].text = str(item['qty'])
                row[1].text = item['part']
                row[2].text = item['desc']
                row[3].text = f"Rp. {int(item['price']):,}".replace(",", ".")
            break

    # Save Word file
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# UI
st.title("Generator Penawaran Otomatis")

# Form
with st.form("penawaran_form"):
    nama_rs = st.text_input("Nama RS")
    alamat = st.text_area("Alamat")
    principal_unit = st.text_input("Nama Principal dan Unit")
    tanggal = st.date_input("Tanggal", datetime.today())
    nomor_urut = st.text_input("Nomor Surat Urut", "993")

    # Item list
    st.markdown("### Item")
    items = []
    for i in range(1, 6):
        with st.expander(f"Item {i}"):
            qty = st.number_input(f"Qty {i}", min_value=1, key=f"qty{i}")
            part = st.text_input(f"Part Number {i}", key=f"part{i}")
            desc = st.text_area(f"Description {i}", key=f"desc{i}")
            price = st.number_input(f"Harga {i}", min_value=0, key=f"price{i}")
            if part:
                items.append({"qty": qty, "part": part, "desc": desc, "price": price})

    sales_dict = {
        "Rully Candra": "0813-1515-4142",
        "Muhammad Lukmansyah": "0821-2291-1020",
        "Denny Firmansyah": "0821-1408-0011",
        "Alamas Ramadhan": "0857-7376-2820"
    }
    sales = st.selectbox("Nama Sales", list(sales_dict.keys()))
    phone = sales_dict[sales]

    submitted = st.form_submit_button("Generate")

# Processing
if submitted:
    bulan = romawi(tanggal.month)
    tahun_kode = str(tanggal.year)[-2:]
    no_surat = f"{nomor_urut}/JKT/SRV/AA/{bulan}/{tahun_kode}"

    subtotal = sum(item['price'] for item in items)
    ppn = int(subtotal * 0.11)
    total = subtotal + ppn

    data = {
        "{{RS}}": nama_rs,
        "{{ALAMAT}}": alamat,
        "{{NOSURAT}}": no_surat,
        "{{TANGGAL}}": tanggal.strftime("%d %B %Y"),
        "{{UNIT}}": principal_unit,
        "{{SUBTOTAL}}": f"Rp. {subtotal:,}".replace(",", "."),
        "{{PPN}}": f"Rp. {ppn:,}".replace(",", "."),
        "{{TOTAL}}": f"Rp. {total:,}".replace(",", "."),
        "{{SALES}}": sales,
        "{{PHONE}}": phone,
    }

    word_file = generate_word(data, items)

    st.download_button("Download Word", word_file, file_name="Penawaran.docx")
