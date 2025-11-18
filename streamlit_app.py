import streamlit as st
from datetime import date
import io
from docx import Document
from docx.shared import Inches, Pt
import os

# Try to import reportlab with fallback installation
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (
        SimpleDocTemplate,
        Table,
        TableStyle,
        Paragraph,
        Spacer,
    )
    from reportlab.lib.units import cm
    from reportlab.lib.pagesizes import A4
except ImportError:
    import subprocess
    import sys

    subprocess.check_call([sys.executable, "-m", "pip", "install", "reportlab"])
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (
        SimpleDocTemplate,
        Table,
        TableStyle,
        Paragraph,
        Spacer,
    )
    from reportlab.lib.units import cm
    from reportlab.lib.pagesizes import A4


def format_rupiah(angka):
    return "Rp. {:,.0f}".format(angka).replace(",", ".")


def format_tanggal_indonesia(tanggal):
    bulan_dict = {
        "January": "Januari",
        "February": "Februari",
        "March": "Maret",
        "April": "April",
        "May": "Mei",
        "June": "Juni",
        "July": "Juli",
        "August": "Agustus",
        "September": "September",
        "October": "Oktober",
        "November": "November",
        "December": "Desember",
    }
    hari = tanggal.day
    bulan = bulan_dict[tanggal.strftime("%B")]
    tahun = tanggal.year
    return f"{hari} {bulan} {tahun}"


def create_pdf(
    nama_customer,
    alamat,
    nomor_penawaran,
    tanggal,
    nama_unit,
    items,
    ketersediaan,
    pic,
    pic_telp,
    subtotal1,
    subtotal2,
    ppn,
    total,
    diskon_option,
    diskon_value,
    selected_items,
    price_diskon,
):
    buffer = io.BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    style_normal = styles["Normal"]
    style_bold = ParagraphStyle(
        "Bold",
        parent=style_normal,
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=13,
    )

    story = []

    # Header
    story.append(Paragraph("Kepada Yth:", style_bold))
    story.append(Paragraph(nama_customer or "-", style_bold))
    story.append(
        Paragraph(
            alamat.replace("\n", "<br/>") if alamat else "-",
            style_normal,
        )
    )
    story.append(Spacer(1, 0.4 * cm))

    story.append(
        Paragraph(f"Nomor Penawaran: {nomor_penawaran or '-'}", style_normal)
    )
    story.append(
        Paragraph(f"Tanggal: {format_tanggal_indonesia(tanggal)}", style_normal)
    )
    story.append(Paragraph(f"Unit: {nama_unit or '-'}", style_normal))
    story.append(Spacer(1, 0.6 * cm))

    intro_text = (
        "Terima kasih atas kesempatan yang telah diberikan kepada kami. "
        "Bersama ini kami mengajukan penawaran harga item sebagai berikut:"
    )
    story.append(Paragraph(intro_text, style_normal))
    story.append(Spacer(1, 0.4 * cm))

    # Tabel item
    data_table = [
        [
            Paragraph("Qty", style_bold),
            Paragraph("Part Number", style_bold),
            Paragraph("Description", style_bold),
            Paragraph("Price per item", style_bold),
            Paragraph("Total Price", style_bold),
        ]
    ]

    for item in items:
        qty_text = f"{item['qty']} {item['uom']}".strip()
        partnumber_text = item["partnumber"] or ""
        description_text = item["description"] or ""
        priceperitem_text = format_rupiah(item["priceperitem"])
        price_text = format_rupiah(item["price"])

        data_table.append(
            [
                Paragraph(qty_text, style_normal),
                Paragraph(partnumber_text, style_normal),
                Paragraph(description_text, style_normal),
                Paragraph(priceperitem_text, style_normal),
                Paragraph(price_text, style_normal),
            ]
        )

    table = Table(
        data_table,
        colWidths=[2 * cm, 3 * cm, 7 * cm, 3 * cm, 3 * cm],
        repeatRows=1,
    )

    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("ALIGN", (0, 0), (-1, 0), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ALIGN", (0, 1), (0, -1), "CENTER"),  # Qty
                ("ALIGN", (3, 1), (-1, -1), "RIGHT"),  # Harga & Total
                ("FONTSIZE", (0, 0), (-1, -1), 9),
            ]
        )
    )

    story.append(table)
    story.append(Spacer(1, 0.4 * cm))

    # Ringkasan harga (Sub Total, Diskon, PPN, Total)
    summary_data = []

    summary_data.append(
        [
            Paragraph("", style_normal),
            Paragraph("", style_normal),
            Paragraph("", style_normal),
            Paragraph("Sub Total I", style_bold),
            Paragraph(format_rupiah(subtotal1), style_normal),
        ]
    )

    if price_diskon > 0:
        if diskon_option == "Diskon persentase (%)":
            label_diskon = f"Diskon {round(diskon_value)}%"
        else:
            label_diskon = "Diskon (Rp)"

        if selected_items:
            label_diskon += " (" + ", ".join([f"Item {i+1}" for i in selected_items]) + ")"

        summary_data.append(
            [
                Paragraph("", style_normal),
                Paragraph("", style_normal),
                Paragraph("", style_normal),
                Paragraph(label_diskon, style_normal),
                Paragraph(f"-{format_rupiah(price_diskon)}", style_normal),
            ]
        )

    summary_data.append(
        [
            Paragraph("", style_normal),
            Paragraph("", style_normal),
            Paragraph("", style_normal),
            Paragraph("Sub Total II", style_bold),
            Paragraph(format_rupiah(subtotal2), style_normal),
        ]
    )

    summary_data.append(
        [
            Paragraph("", style_normal),
            Paragraph("", style_normal),
            Paragraph("", style_normal),
            Paragraph("PPN 11%", style_bold),
            Paragraph(format_rupiah(ppn), style_normal),
        ]
    )

    summary_data.append(
        [
            Paragraph("", style_normal),
            Paragraph("", style_normal),
            Paragraph("", style_normal),
            Paragraph("TOTAL", style_bold),
            Paragraph(format_rupiah(total), style_bold),
        ]
    )

    summary_table = Table(
        summary_data,
        colWidths=[2 * cm, 3 * cm, 7 * cm, 3 * cm, 3 * cm],
    )

    summary_table.setStyle(
        TableStyle(
            [
                ("GRID", (3, 0), (-1, -1), 0.5, colors.black),
                ("ALIGN", (3, 0), (-1, -1), "RIGHT"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
            ]
        )
    )

    story.append(summary_table)
    story.append(Spacer(1, 0.6 * cm))

    # Ketersediaan & PIC
    if ketersediaan and ketersediaan != "Jangan tampilkan":
        story.append(
            Paragraph(f"Ketersediaan Barang: {ketersediaan}", style_normal)
        )
    story.append(
        Paragraph(f"PIC: {pic} - {pic_telp}", style_normal)
    )

    story.append(Spacer(1, 0.6 * cm))

    story.append(Paragraph("Hormat kami,", style_normal))
    story.append(Paragraph("PT. IDS Medical Systems Indonesia", style_normal))
    story.append(Spacer(1, 1.5 * cm))
    story.append(Paragraph("M. Athur Yassin", style_normal))
    story.append(Paragraph("Manager II - Engineering", style_normal))

    doc.build(story)
    buffer.seek(0)
    return buffer


# =============================
#   STATE & CHAT-BASED INPUT
# =============================
st.markdown(
    "<h1 style='text-align: center;'>Penawaran Harga (Mode Chat)</h1>",
    unsafe_allow_html=True,
)

# Inisialisasi session_state
if "parsed_data" not in st.session_state:
    st.session_state.parsed_data = {
        "nama_customer": "",
        "alamat": "",
        "nomor_penawaran": "",
        "tanggal": date.today(),
        "nama_unit": "",
        "items": [],
        "diskon_option": "Tanpa diskon",
        "diskon_value": 0,
        "selected_items": [],
        "ketersediaan": "Jangan tampilkan",
        "pic": "Muhammad Lukmansyah",
    }

pic_options = {
    "Muhammad Lukmansyah": "0821 2291 1020",
    "Rully Candra": "0813 1515 4142",
    "Denny Firmansyah": "0821 1408 0011",
    "Alamas Ramadhan": "0857 7376 2820",
}

st.write("Tulis instruksi penawaran di bawah ini (seperti ngobrol):")
st.code(
    """Contoh:
Buat penawaran untuk:
Customer: RS Harapan Sejahtera
Alamat: Jl. Merdeka No. 10, Jakarta
Nomor penawaran: 023/PNW/RS-HS/2025
Tanggal: 18-11-2025
Unit: Ventilator XYZ SN12345

Item:
1) Qty 2 unit, Part Number VT-100, Description Ventilator tipe 100, Harga 45.000.000
2) Qty 3 set, Part Number TS-01, Description Trolley Stand, Harga 5.000.000

Diskon: 10% untuk semua item
Ketersediaan: Ready stock
PIC: Muhammad Lukmansyah
""",
    language="text",
)

chat_input = st.text_area(
    "Chat Penawaran", height=250, help="Tulis detail penawaran secara bebas."
)

if st.button("💬 Proses Chat"):
    import re

    text = chat_input

    def get_after(label):
        for line in text.splitlines():
            if label.lower() in line.lower():
                return line.split(":", 1)[-1].strip()
        return ""

    # Field dasar
    nama_customer = get_after("Customer")
    alamat = get_after("Alamat")
    nomor_penawaran = get_after("Nomor penawaran")
    unit = get_after("Unit")
    tgl_str = get_after("Tanggal")

    # Tanggal dd-mm-yyyy atau dd/mm/yyyy
    parsed_date = date.today()
    if tgl_str:
        match = re.search(r"(\d{1,2})[-/](\d{1,2})[-/](\d{4})", tgl_str)
        if match:
            d, m, y = map(int, match.groups())
            try:
                parsed_date = date(y, m, d)
            except Exception:
                parsed_date = date.today()

    # Diskon
    diskon_option = "Tanpa diskon"
    diskon_value = 0
    selected_items = []

    for line in text.splitlines():
        if "diskon" in line.lower():
            angka = re.findall(r"\d+", line)
            if angka:
                val = int(angka[0])
                if "%" in line:
                    diskon_option = "Diskon persentase (%)"
                    diskon_value = val
                else:
                    diskon_option = "Diskon nominal (Rp)"
                    diskon_value = val
            if "semua" in line.lower():
                selected_items = "semua"
            break

    # Ketersediaan
    ketersediaan = "Jangan tampilkan"
    for line in text.splitlines():
        lower_line = line.lower()
        if (
            "ketersediaan" in lower_line
            or "ready stock" in lower_line
            or "indent" in lower_line
        ):
            if "ready stock" in lower_line:
                ketersediaan = "Ready stock"
            elif "indent" in lower_line:
                ketersediaan = "Indent"
            elif "persediaan masih ada" in lower_line:
                ketersediaan = "Ready jika persediaan masih ada"
            break

    # PIC
    pic = "Muhammad Lukmansyah"
    for name in pic_options.keys():
        if name.lower() in text.lower():
            pic = name
            break

    # Item
    items = []
    for line in text.splitlines():
        # 1) Qty 2 unit, Part Number VT-100, Description ..., Harga 45.000.000
        if re.match(r"\s*\d+\)", line.strip()):
            qty_match = re.search(r"qty\s+(\d+)", line, re.IGNORECASE)
            uom_match = re.search(r"qty\s+\d+\s+(\w+)", line, re.IGNORECASE)
            pn_match = re.search(
                r"part number\s*([A-Za-z0-9\-\_/]+)", line, re.IGNORECASE
            )
            desc_match = re.search(
                r"description\s*(.*?),\s*harga", line, re.IGNORECASE
            )
            harga_match = re.search(r"harga\s*([\d\.]+)", line, re.IGNORECASE)

            qty = qty_match.group(1) if qty_match else "1"
            uom = uom_match.group(1) if uom_match else ""
            partnumber = pn_match.group(1) if pn_match else ""
            description = desc_match.group(1) if desc_match else ""
            priceperitem = 0
            if harga_match:
                h_str = harga_match.group(1).replace(".", "")
                try:
                    priceperitem = int(h_str)
                except Exception:
                    priceperitem = 0

            total = float(qty) * priceperitem
            items.append(
                {
                    "qty": qty,
                    "uom": uom,
                    "partnumber": partnumber,
                    "description": description,
                    "priceperitem": priceperitem,
                    "price": total,
                }
            )

    # Diskon utk semua item
    if selected_items == "semua":
        selected_items = list(range(len(items)))
    elif isinstance(selected_items, str):
        selected_items = []

    # Simpan ke session_state
    st.session_state.parsed_data = {
        "nama_customer": nama_customer,
        "alamat": alamat,
        "nomor_penawaran": nomor_penawaran,
        "tanggal": parsed_date,
        "nama_unit": unit,
        "items": items,
        "diskon_option": diskon_option,
        "diskon_value": diskon_value,
        "selected_items": selected_items,
        "ketersediaan": ketersediaan,
        "pic": pic,
    }

    st.success(
        "Chat berhasil diproses. Silakan cek dan koreksi data di bawah sebelum generate dokumen."
    )

# =============================
#  REVIEW & GENERATE DOKUMEN
# =============================
data = st.session_state.parsed_data

st.markdown("### Data Penawaran (hasil dari chat, bisa dikoreksi)")

col_a, col_b = st.columns(2)
with col_a:
    data["nama_customer"] = st.text_input(
        "Nama Customer", value=data["nama_customer"]
    )
    data["alamat"] = st.text_area("Alamat Customer", value=data["alamat"])
    data["nomor_penawaran"] = st.text_input(
        "Nomor Penawaran", value=data["nomor_penawaran"]
    )
with col_b:
    data["tanggal"] = st.date_input("Tanggal", value=data["tanggal"])
    data["nama_unit"] = st.text_input(
        "Nama Unit (Tipe dan Serial Number jika ada)", value=data["nama_unit"]
    )

st.markdown("#### Item yang ditawarkan")
for i, item in enumerate(data["items"]):
    st.write(f"**Item {i+1}**")
    st.text(
        f"Qty: {item['qty']} {item['uom']}\n"
        f"Part Number: {item['partnumber']}\n"
        f"Description: {item['description']}\n"
        f"Harga per item: {format_rupiah(item['priceperitem'])}\n"
        f"Total: {format_rupiah(item['price'])}"
    )

# Diskon
st.markdown("#### Diskon")
diskon_option = st.radio(
    "Jenis Diskon",
    ["Tanpa diskon", "Diskon persentase (%)", "Diskon nominal (Rp)"],
    index=[
        "Tanpa diskon",
        "Diskon persentase (%)",
        "Diskon nominal (Rp)",
    ].index(data["diskon_option"]),
)
diskon_value = 0
selected_items = []

jumlah_item = len(data["items"])

if diskon_option != "Tanpa diskon" and jumlah_item > 0:
    if jumlah_item > 1:
        diskon_scope = st.radio(
            "Diskon berlaku untuk:", ["Semua item", "Pilih item tertentu"], index=0
        )
        if diskon_scope == "Pilih item tertentu":
            st.markdown("**Pilih item yang dapat diskon:**")
            cols = st.columns(3)
            for i in range(jumlah_item):
                with cols[i % 3]:
                    if st.checkbox(f"Item {i+1}", key=f"diskon_item_{i}"):
                        selected_items.append(i)
        else:
            selected_items = list(range(jumlah_item))
    else:
        selected_items = [0]

    if diskon_option == "Diskon persentase (%)":
        diskon_value = st.number_input(
            "Besar diskon (%)",
            min_value=0,
            max_value=100,
            value=data["diskon_value"],
            format="%d",
        )
    else:
        diskon_value = st.number_input(
            "Besar diskon (Rp)",
            min_value=0,
            value=data["diskon_value"],
            format="%d",
        )
else:
    diskon_value = 0
    selected_items = []

data["diskon_option"] = diskon_option
data["diskon_value"] = diskon_value
data["selected_items"] = selected_items

# Ketersediaan & PIC
opsi_ketersediaan = [
    "Jangan tampilkan",
    "Ready stock",
    "Ready jika persediaan masih ada",
    "Indent",
]
data["ketersediaan"] = st.selectbox(
    "Ketersediaan Barang",
    opsi_ketersediaan,
    index=opsi_ketersediaan.index(data["ketersediaan"])
    if data["ketersediaan"] in opsi_ketersediaan
    else 0,
)
data["pic"] = st.selectbox(
    "Nama PIC",
    list(pic_options.keys()),
    index=list(pic_options.keys()).index(data["pic"])
    if data["pic"] in pic_options
    else 0,
)
pic_telp = pic_options[data["pic"]]

# =============================
#  GENERATE DOKUMEN
# =============================
if st.button("📥 Generate Dokumen Penawaran"):
    items = data["items"]
    nama_customer = data["nama_customer"]
    alamat = data["alamat"]
    nomor_penawaran = data["nomor_penawaran"]
    tanggal = data["tanggal"]
    nama_unit = data["nama_unit"]
    diskon_option = data["diskon_option"]
    diskon_value = data["diskon_value"]
    selected_items = data["selected_items"]
    ketersediaan = data["ketersediaan"]
    pic = data["pic"]

    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"

    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]

    # Ganti path sesuai lokasi kop surat Anda
    image_path = "/mnt/data/92e028fb-3499-479f-a167-62ec17940b2d.png"
    if os.path.exists(image_path):
        try:
            header_para.add_run().add_picture(image_path, width=Inches(6.5))
        except Exception as e:
            st.warning(f"Gagal menambahkan kop surat: {e}")

    deskripsi_item = (
        items[0]["description"][:30].replace("/", "-")
        if items and items[0]["description"]
        else "Penawaran"
    )
    safe_customer = (nama_customer or "Customer").replace(" ", "_")
    safe_unit = (nama_unit or "Unit").replace(" ", "_")
    nama_file = (
        f"{nomor_penawaran}_{safe_customer}-{safe_unit}_"
        f"{deskripsi_item.replace(' ', '_')}.docx"
    )

    # Konten dokumen
    p = doc.add_paragraph()
    run = p.add_run("Kepada Yth: ")
    run.bold = True
    run = p.add_run(nama_customer)
    run.bold = True
    doc.add_paragraph(alamat)

    p = doc.add_paragraph()
    run = p.add_run("Hal: Penawaran Harga")
    run.bold = True
    run.underline = True
    p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    run = p.add_run(f"No: {nomor_penawaran}/JKT/SRV/AA/25")
    run.add_tab()
    run.add_tab()
    run.add_tab()
    run.add_tab()
    run.add_tab()
    run.add_text(f"Jakarta, {format_tanggal_indonesia(tanggal)}")
    p.paragraph_format.space_after = Pt(0)

    doc.add_paragraph()
    p = doc.add_paragraph(
        f"Terima kasih atas kesempatan yang telah diberikan kepada kami. "
        f"Bersama ini kami mengajukan penawaran harga item untuk unit {nama_unit} "
        f"di {nama_customer}, sebagai berikut:\n"
    )
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Qty"
    hdr_cells[1].text = "Part Number"
    hdr_cells[2].text = "Description"
    hdr_cells[3].text = "Price per item"
    hdr_cells[4].text = "Total Price"

    table.columns[0].width = Inches(1)
    table.columns[1].width = Inches(1)
    table.columns[2].width = Inches(2)
    table.columns[3].width = Inches(1)
    table.columns[4].width = Inches(1)

    subtotal1 = 0
    for i, item in enumerate(items):
        row_cells = table.add_row().cells
        row_cells[0].text = f"{item['qty']} {item['uom']}"
        row_cells[1].text = item["partnumber"]
        row_cells[2].text = item["description"]
        row_cells[3].text = format_rupiah(item["priceperitem"])
        row_cells[4].text = format_rupiah(item["price"])
        subtotal1 += item["price"]

        for cell in row_cells:
            cell.paragraphs[0].alignment = 1

    # Hitung diskon
    price_diskon = 0
    if diskon_option != "Tanpa diskon" and selected_items:
        if diskon_option == "Diskon persentase (%)":
            for i in selected_items:
                price_diskon += items[i]["price"] * (diskon_value / 100)
        else:
            total_terdiskon = sum(items[i]["price"] for i in selected_items)
            if total_terdiskon > 0:
                for i in selected_items:
                    price_diskon += (
                        items[i]["price"] / total_terdiskon
                    ) * diskon_value
        price_diskon = round(price_diskon)

    subtotal2 = subtotal1 - price_diskon
    ppn = subtotal2 * 0.11
    total = subtotal2 + ppn

    for label, value in [
        ("Sub Total I", subtotal1),
        ("Sub Total II", subtotal2),
        ("PPN 11%", ppn),
        ("TOTAL", total),
    ]:
        row = table.add_row().cells
        row[3].text = label
        row[4].text = format_rupiah(value)
        for cell in row:
            cell.paragraphs[0].alignment = 1
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    if price_diskon > 0:
        row_disc = table.add_row().cells
        if diskon_option == "Diskon persentase (%)":
            row_disc[3].text = (
                f"Diskon {round(diskon_value)}% ("
                f"{', '.join(['Item ' + str(i+1) for i in selected_items])})"
            )
        else:
            row_disc[3].text = (
                f"Diskon (Rp) ("
                f"{', '.join(['Item ' + str(i+1) for i in selected_items])})"
            )
        row_disc[4].text = f"-{format_rupiah(price_diskon)}"
        for cell in row_disc:
            cell.paragraphs[0].alignment = 1
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    for text in [
        "\nSyarat dan ketentuan:",
        "Harga                               : Sudah termasuk PPN 11%",
        "Pembayaran                   : Tunai atau transfer",
        "Masa berlaku                 : 2 minggu",
    ]:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(0)

    if ketersediaan != "Jangan tampilkan":
        p = doc.add_paragraph(f"Ketersediaan Barang    : {ketersediaan}")
        p.paragraph_format.space_after = Pt(0)

    for text in ["\nHormat kami,", "PT. IDS Medical Systems Indonesia"]:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(0)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    for text in ["M. Athur Yassin", "Manager II - Engineering", pic, pic_telp]:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(0)

    # Save dokumen Word
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Create PDF (pakai layout rapi)
    pdf_buffer = create_pdf(
        nama_customer,
        alamat,
        nomor_penawaran,
        tanggal,
        nama_unit,
        items,
        ketersediaan,
        pic,
        pic_telp,
        subtotal1,
        subtotal2,
        ppn,
        total,
        diskon_option,
        diskon_value,
        selected_items,
        price_diskon,
    )

    # Preview
    preview_doc = Document(buffer)
    preview_text = "\n".join([para.text for para in preview_doc.paragraphs])

    st.markdown("### Preview Penawaran")
    st.text_area("Isi Penawaran", value=preview_text, height=400)

    # Download buttons
    st.download_button(
        label="⬇️ Download Penawaran (Word)",
        data=buffer,
        file_name=nama_file,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    st.download_button(
        label="⬇️ Download Penawaran (PDF)",
        data=pdf_buffer,
        file_name=nama_file.replace(".docx", ".pdf"),
        mime="application/pdf",
    )
